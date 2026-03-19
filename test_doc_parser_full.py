#!/usr/bin/env python
"""Test DocumentParser with sample DOCX file"""

import asyncio
import os
import tempfile
from docx import Document

async def test_document_parser():
    """Test the DocumentParser with a real DOCX file"""
    
    print("=" * 60)
    print("Testing DocumentParser Module")
    print("=" * 60)
    
    # Test 1: Import
    print("\n[Test 1] Importing DocumentParser...")
    try:
        from backend.doc_parser import DocumentParser
        print("✓ DocumentParser imported successfully")
    except ImportError as e:
        print(f"✗ Import failed: {e}")
        return False
    
    # Test 2: Instantiation
    print("\n[Test 2] Creating parser instance...")
    try:
        parser = DocumentParser()
        print("✓ DocumentParser instantiated")
    except Exception as e:
        print(f"✗ Instantiation failed: {e}")
        return False
    
    # Test 3: Create a sample DOCX
    print("\n[Test 3] Creating sample DOCX file...")
    try:
        doc = Document()
        doc.add_heading('John Smith - Professional Profile', 0)
        doc.add_paragraph('Senior Software Engineer with 8 years of experience')
        doc.add_heading('Skills', level=1)
        doc.add_paragraph('Python, JavaScript, FastAPI, React, PostgreSQL')
        doc.add_heading('Services', level=1)
        doc.add_paragraph('- Custom software development')
        doc.add_paragraph('- API design and optimization')
        doc.add_paragraph('- Team leadership and mentoring')
        doc.add_paragraph('- Cloud infrastructure consulting')
        doc.add_heading('Experience', level=1)
        doc.add_paragraph('Tech Lead at TechCorp (5 years) - Led development of 3 major platforms')
        doc.add_paragraph('Senior Developer at StartupXYZ (3 years) - Built full-stack applications')
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            temp_docx = tmp.name
        
        print(f"✓ Sample DOCX created: {temp_docx}")
    except Exception as e:
        print(f"✗ DOCX creation failed: {e}")
        return False
    
    # Test 4: Extract text
    print("\n[Test 4] Extracting text from DOCX...")
    try:
        text = parser.extract_text(temp_docx, ".docx")
        print(f"✓ Text extracted ({len(text)} characters)")
        print(f"  First 100 chars: {text[:100]}...")
    except Exception as e:
        print(f"✗ Text extraction failed: {e}")
        return False
    finally:
        if os.path.exists(temp_docx):
            os.remove(temp_docx)
    
    # Test 5: Fallback response (when ASI-1 not configured)
    print("\n[Test 5] Testing fallback response...")
    try:
        fallback = parser._create_fallback_response()
        assert "person_or_company" in fallback
        assert "services" in fallback
        assert isinstance(fallback["services"], list)
        print(f"✓ Fallback response valid")
        print(f"  - Keys: {list(fallback.keys())}")
        print(f"  - Services array: {fallback['services']}")
    except Exception as e:
        print(f"✗ Fallback response test failed: {e}")
        return False
    
    # Test 6: Main API endpoints check
    print("\n[Test 6] Checking API endpoints...")
    try:
        from main import app, DocumentAnalysisRequest, doc_parser
        routes = [str(r.path) for r in app.routes]
        
        required_routes = [
            "/api/documents/analyze",
            "/api/documents/supported-formats"
        ]
        
        for route in required_routes:
            if any(route in r for r in routes):
                print(f"✓ {route} endpoint found")
            else:
                print(f"✗ {route} endpoint NOT found")
                return False
    except Exception as e:
        print(f"✗ API endpoint check failed: {e}")
        return False
    
    print("\n" + "=" * 60)
    print("✓ All tests passed!")
    print("=" * 60)
    return True

if __name__ == "__main__":
    success = asyncio.run(test_document_parser())
    exit(0 if success else 1)
